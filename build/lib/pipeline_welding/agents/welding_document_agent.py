from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


@dataclass(frozen=True)
class WeldingDocumentAgentConfig:
    template_docx_path: Path = Path("data/MHPWPS-062.docx")
    output_dir: Path = Path("result")
    output_prefix: str = "welding_standard_filled"


class WeldingDocumentAgent:
    """Fills a WPS-style DOCX from the standard-agent JSON result."""

    def __init__(self, config: WeldingDocumentAgentConfig | None = None) -> None:
        self.config = config or WeldingDocumentAgentConfig()

    def build_document(self, standard_result: dict[str, Any]) -> Path:
        template_path = self.config.template_docx_path
        if not template_path.exists():
            raise FileNotFoundError(f"Template DOCX not found: {template_path}")

        self.config.output_dir.mkdir(parents=True, exist_ok=True)
        document = Document(str(template_path))

        document_fields = self._extract_document_fields(standard_result)
        document_fields = self._ensure_required_document_fields(document_fields)
        self._fill_known_paragraphs(document, document_fields)
        self._fill_known_tables(document, document_fields)

        output_path = self._build_output_path()
        document.save(str(output_path))
        print(f"已生成焊接标准文档：{output_path}")
        print("已写入字段：")
        for key, value in document_fields.items():
            print(f"  {key}: {value}")
        return output_path

    def _extract_document_fields(self, standard_result: dict[str, Any]) -> dict[str, str]:
        document_fields = standard_result.get("document_fields", {})
        if isinstance(document_fields, dict) and document_fields:
            return {key: self._clean_value(value) for key, value in document_fields.items()}

        fields = standard_result.get("input", {})
        if not isinstance(fields, dict):
            return {}
        fallback = {key: self._clean_value(value) for key, value in fields.items()}
        fallback.update(self._extract_reference_values(standard_result))
        return fallback

    def _fill_known_paragraphs(
        self,
        document: Document,
        document_fields: dict[str, str],
    ) -> None:
        paragraphs = self._iter_all_paragraphs(document)
        self._fill_first_label(paragraphs, "单位名称", document_fields, "unit_name")
        self._fill_first_label(paragraphs, "预焊接工艺规程编号", document_fields, "wps_no")
        self._fill_first_label(paragraphs, "所支持的焊接工艺评定编号", document_fields, "pqr_no")
        self._fill_first_label(paragraphs, "焊接方法", document_fields, "welding_process")
        self._fill_first_label(paragraphs, "机械化程度", document_fields, "mechanization")
        self._fill_first_label(paragraphs, "类别号", document_fields, "base_material_category")
        self._fill_label_after(paragraphs, "组别号", self._value(document_fields, "base_material_group"))
        self._fill_label_after(paragraphs, "相焊或标准号", self._value(document_fields, "base_material_standard"))
        self._fill_label_after(paragraphs, "材料代号", self._value(document_fields, "base_material_grade"))
        self._fill_first_label(paragraphs, "对接焊缝焊件母材厚度范围", document_fields, "base_thickness_range_butt")
        self._fill_first_label(paragraphs, "角焊缝焊件母材厚度范围", document_fields, "base_thickness_range_fillet")
        self._fill_label_after(
            paragraphs,
            "管子直径、壁厚范围：对接焊缝",
            self._value(document_fields, "pipe_diameter_thickness_butt"),
        )
        self._fill_label_after_with_anchor(
            paragraphs,
            "管子直径、壁厚范围",
            "角焊缝",
            self._value(document_fields, "pipe_diameter_thickness_fillet"),
        )
        self._fill_first_label(paragraphs, "对接焊缝焊件焊缝金属厚度范围", document_fields, "base_thickness_range_butt")
        self._fill_first_label(paragraphs, "角焊缝焊件焊缝金属厚度范围", document_fields, "base_thickness_range_fillet")
        self._fill_label_after(paragraphs, "其他：", self._value(document_fields, "corrosion_overlay_other"))
        self._fill_electrical_paragraphs(paragraphs, document_fields)
        self._fill_technical_paragraphs(paragraphs, document_fields)
        self._fill_signature_paragraphs(paragraphs, document_fields)

    def _fill_known_tables(
        self,
        document: Document,
        document_fields: dict[str, str],
    ) -> None:
        for index, table in enumerate(document.tables):
            if index == 0:
                self._fill_joint_table(table, document_fields)
            if index == 1:
                self._fill_filler_metal_table(table, document_fields)
            if index == 3:
                self._fill_process_condition_table(table, document_fields)
            if self._is_welding_parameter_table(table):
                self._fill_welding_parameter_table(table, document_fields)

    def _fill_joint_table(self, table: Any, document_fields: dict[str, str]) -> None:
        if not table.rows or not table.rows[0].cells:
            return
        paragraphs = table.rows[0].cells[0].paragraphs
        self._fill_label_after(paragraphs, "坡口形式:", self._value(document_fields, "groove_type"))
        self._fill_label_after(paragraphs, "衬垫(材料及规格", self._value(document_fields, "backing"))
        self._fill_label_after(paragraphs, "其他", self._value(document_fields, "joint_other"))

    def _fill_filler_metal_table(self, table: Any, document_fields: dict[str, str]) -> None:
        row_values = {
            "焊材标准": self._value(document_fields, "filler_standard"),
            "填充金属尺寸": self._value(document_fields, "filler_diameter"),
            "焊材型号": self._value(document_fields, "filler_model"),
            "焊材牌号": self._value(document_fields, "filler_trade_name"),
            "填充金属类别": self._value(document_fields, "filler_class"),
            "其他": self._value(document_fields, "filler_category"),
        }
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().replace("：", "")
            for key, value in row_values.items():
                if key in label:
                    row.cells[1].text = value
                    if len(row.cells) > 2:
                        row.cells[2].text = "/"
                    break

    def _fill_process_condition_table(self, table: Any, document_fields: dict[str, str]) -> None:
        if not table.rows:
            return
        cells = table.rows[0].cells
        if len(cells) < 3:
            return
        self._fill_label_after(cells[0].paragraphs, "对接焊缝的位置", self._value(document_fields, "butt_weld_position"))
        self._fill_label_after(cells[0].paragraphs, "立焊的焊接方向：（向上、向下）", self._value(document_fields, "vertical_direction"))
        self._fill_label_after(cells[0].paragraphs, "角焊缝位置", self._value(document_fields, "fillet_weld_position"))
        self._fill_label_after(cells[0].paragraphs, "保温温度（℃）", self._value(document_fields, "pwht_temperature"))
        self._fill_label_after(cells[0].paragraphs, "保温时间范围（h）", self._value(document_fields, "pwht_time"))
        self._fill_label_after(cells[1].paragraphs, "最小预热温度（℃）", self._value(document_fields, "preheat_temperature"))
        self._fill_label_after(cells[1].paragraphs, "最大道间温度（℃）", self._value(document_fields, "interpass_temperature"))
        self._fill_label_after(cells[1].paragraphs, "保持预热时间", self._value(document_fields, "preheat_time"))
        self._fill_label_after(cells[1].paragraphs, "加热方式", self._value(document_fields, "heating_method"))
        self._fill_gas_cell(cells[2], document_fields)

    @staticmethod
    def _is_welding_parameter_table(table: Any) -> bool:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        return all(keyword in text for keyword in ("焊道", "焊接", "电流"))

    def _fill_welding_parameter_table(
        self,
        table: Any,
        document_fields: dict[str, str],
    ) -> None:
        filled = 0
        for row in table.rows:
            cells = row.cells
            if len(cells) < 9 or self._is_parameter_header_row(row):
                continue
            filled += 1
            if filled > 2:
                break
            prefix = f"bead_{filled}_"
            self._set_cell_text(cells, 0, str(filled))
            self._set_cell_text(cells, 1, self._value(document_fields, prefix + "process"))
            self._set_cell_text(cells, 2, self._value(document_fields, prefix + "filler_metal"))
            self._set_cell_text(cells, 3, self._value(document_fields, prefix + "diameter"))
            self._set_cell_text(cells, 4, self._value(document_fields, prefix + "polarity"))
            self._set_cell_text(cells, 5, self._value(document_fields, prefix + "current"))
            self._set_cell_text(cells, 6, self._value(document_fields, prefix + "voltage"))
            self._set_cell_text(cells, 7, self._value(document_fields, prefix + "speed"))
            self._set_cell_text(cells, 8, self._value(document_fields, prefix + "heat_input"))

    @staticmethod
    def _is_parameter_header_row(row: Any) -> bool:
        text = "\n".join(cell.text for cell in row.cells)
        return any(
            keyword in text
            for keyword in (
                "焊道/",
                "焊接\n方法",
                "填充金属",
                "焊接电流",
                "电弧电压",
                "牌号",
                "直径",
                "极性",
                "电流(A)",
            )
        )

    def _extract_reference_values(self, standard_result: dict[str, Any]) -> dict[str, str]:
        text = self._collect_clean_reference_text(standard_result)
        return {
            "preheat_temperature": self._first_match(text, (r"预热温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)",)),
            "interpass_temperature": self._first_match(text, (r"层间温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)", r"道间温度[^\d≥≤]*([≥≤]?\s*\d+\s*℃?)")),
            "current": self._first_match(text, (r"电流[^\d]*(\d+\s*[-~～]\s*\d+\s*A?)",)),
            "voltage": self._first_match(text, (r"电压[^\d]*(\d+\s*[-~～]\s*\d+\s*V?)",)),
            "welding_speed": self._first_match(text, (r"焊接速度[^\d]*(\d+\s*[-~～]\s*\d+\s*(?:mm/min)?)",)),
            "heat_input": self._first_match(text, (r"线能量[^\d]*(\d+(?:\.\d+)?\s*(?:kJ/cm)?)",)),
            "filler_metal": self._first_match(text, (r"(E\d{3,4}[A-Z0-9-]*)",)),
            "filler_diameter": self._first_match(text, (r"(?:φ|Φ)\s*(\d+(?:\.\d+)?\s*mm)",)),
            "polarity": self._first_match(text, (r"(反接|正接|DCEN|DCEP|EP|EN)",)),
        }

    def _collect_clean_reference_text(self, standard_result: dict[str, Any]) -> str:
        parts: list[str] = []
        standard = standard_result.get("pipeline_welding_standard", {})
        if isinstance(standard, dict):
            parts.extend(str(item) for item in standard.get("required_controls", []))

        results = standard_result.get("mcp_search", {}).get("results", {})
        if isinstance(results, dict):
            for items in results.values():
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            parts.extend(
                                str(item.get(key, "")) for key in ("title", "snippet") if item.get(key)
                            )

        return "\n".join(part for part in parts if self._is_clean_text(part))

    @staticmethod
    def _is_clean_text(text: str) -> bool:
        lowered = text.lower()
        blocked = ("not found", "unknown tool", "unknown too", "error", "missing")
        return bool(text.strip()) and not any(item in lowered for item in blocked)

    @staticmethod
    def _first_match(text: str, patterns: tuple[str, ...]) -> str:
        import re

        for pattern in patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return match.group(1).replace(" ", "")
        return ""

    def _fill_technical_paragraphs(
        self,
        paragraphs: Any,
        document_fields: dict[str, str],
    ) -> None:
        self._fill_label_after(paragraphs, "摆动焊或不摆动焊", self._value(document_fields, "weaving"))
        self._fill_label_after(paragraphs, "摆动参数", self._value(document_fields, "weaving_parameter"))
        self._fill_label_after(paragraphs, "焊前清理和层间清理", self._value(document_fields, "cleaning"))
        self._fill_label_after(paragraphs, "背面清根方法", self._value(document_fields, "back_gouging"))
        self._fill_label_after(paragraphs, "单道焊或多道焊（每面）", self._value(document_fields, "single_or_multi_pass"))
        self._fill_label_after(paragraphs, "单丝焊或多丝焊", self._value(document_fields, "single_or_multi_wire"))
        self._fill_label_after(paragraphs, "导电嘴至工件距离（mm）", self._value(document_fields, "contact_tip_distance"))
        self._fill_label_after(paragraphs, "锤击", self._value(document_fields, "peening"))
        self._fill_label_after(paragraphs, "其他：", self._value(document_fields, "technical_other"))

    def _fill_signature_paragraphs(
        self,
        paragraphs: Any,
        document_fields: dict[str, str],
    ) -> None:
        self._fill_label_after(paragraphs, "编制", self._value(document_fields, "prepared_by"))
        self._fill_label_after(paragraphs, "审核", self._value(document_fields, "reviewed_by"))
        self._fill_label_after(paragraphs, "批准", self._value(document_fields, "approved_by"))

    def _fill_electrical_paragraphs(
        self,
        paragraphs: Any,
        document_fields: dict[str, str],
    ) -> None:
        self._fill_label_after(paragraphs, "电流种类", self._value(document_fields, "current_type"))
        self._fill_label_after(paragraphs, "极性", self._value(document_fields, "polarity"))
        self._fill_label_after(paragraphs, "焊接电流范围（A）", self._value(document_fields, "current"))
        self._fill_label_after(paragraphs, "电弧电压（V）", self._value(document_fields, "voltage"))
        self._fill_label_after(paragraphs, "焊接速度（范围）", self._value(document_fields, "welding_speed"))
        self._fill_label_after(paragraphs, "钨极类型及直径", self._value(document_fields, "tungsten_electrode"))
        self._fill_label_after(paragraphs, "喷嘴直径（mm）", self._value(document_fields, "nozzle_diameter"))
        self._fill_label_after(paragraphs, "焊接电弧种类（喷射弧、短路弧等）", self._value(document_fields, "arc_type"))
        self._fill_label_after(paragraphs, "焊丝送进速度（cm/min）", self._value(document_fields, "wire_feed_speed"))

    def _iter_all_paragraphs(self, document: Document) -> list[Any]:
        paragraphs = list(document.paragraphs)
        for table in document.tables:
            paragraphs.extend(self._iter_table_paragraphs(table))
        return paragraphs

    def _iter_table_paragraphs(self, table: Any) -> list[Any]:
        paragraphs: list[Any] = []
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                paragraphs.extend(cell.paragraphs)
                for nested_table in cell.tables:
                    paragraphs.extend(self._iter_table_paragraphs(nested_table))
        return paragraphs

    def _fill_gas_cell(self, cell: Any, document_fields: dict[str, str]) -> None:
        values = (
            f"{self._value(document_fields, 'shielding_gas')} "
            f"{self._value(document_fields, 'shielding_gas_mix')} "
            f"{self._value(document_fields, 'gas_flow')}"
        )
        self._fill_label_after(cell.paragraphs, "保护气", values)
        self._fill_label_after(cell.paragraphs, "尾部保护气", self._value(document_fields, "trailing_gas"))
        self._fill_label_after(cell.paragraphs, "背面保护气", self._value(document_fields, "backing_gas"))

    def _fill_first_label(
        self,
        paragraphs: Any,
        label: str,
        document_fields: dict[str, str],
        key: str,
    ) -> None:
        self._fill_label_after(paragraphs, label, self._value(document_fields, key))

    def _build_output_path(self) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return self.config.output_dir / f"{self.config.output_prefix}_{timestamp}.docx"

    @staticmethod
    def _fill_label_after(paragraphs: Any, keyword: str, value: str) -> None:
        for paragraph in paragraphs:
            if keyword in paragraph.text:
                WeldingDocumentAgent._write_after_keyword(paragraph, keyword, value)
                return

    @staticmethod
    def _fill_label_after_with_anchor(paragraphs: Any, anchor: str, keyword: str, value: str) -> None:
        for paragraph in paragraphs:
            if anchor in paragraph.text and keyword in paragraph.text:
                WeldingDocumentAgent._write_after_keyword(paragraph, keyword, value)
                return

    @staticmethod
    def _write_after_keyword(paragraph: Any, keyword: str, value: str) -> None:
        clean_value = WeldingDocumentAgent._clean_value(value)
        saw_keyword = False
        for run in paragraph.runs:
            text = run.text or ""
            if saw_keyword and (not text.strip() or text.strip() == "/"):
                run.text = clean_value
                return
            if keyword in text:
                saw_keyword = True
                suffix = text.split(keyword, 1)[1].strip()
                if suffix and suffix != "/" and len(suffix) <= 18:
                    run.text = text.split(keyword, 1)[0] + keyword + clean_value
                    return
                if suffix and suffix != "/":
                    continue

        if paragraph.runs:
            paragraph.runs[-1].text = f"{paragraph.runs[-1].text}{clean_value}"
        else:
            paragraph.add_run(clean_value)

    @staticmethod
    def _set_cell_text(cells: Any, index: int, value: str | None) -> None:
        if index < len(cells):
            cells[index].text = WeldingDocumentAgent._clean_value(value)

    @staticmethod
    def _ensure_required_document_fields(document_fields: dict[str, str]) -> dict[str, str]:
        required_keys = (
            "unit_name",
            "welding_process",
            "welding_object",
            "joint_type",
            "base_material",
            "base_thickness_or_diameter",
            "wps_no",
            "pqr_no",
            "mechanization",
            "groove_type",
            "backing",
            "joint_other",
            "base_material_category",
            "base_material_group",
            "base_material_standard",
            "base_material_grade",
            "base_thickness_range_butt",
            "base_thickness_range_fillet",
            "pipe_diameter_thickness_butt",
            "pipe_diameter_thickness_fillet",
            "corrosion_overlay_chemical_composition",
            "corrosion_overlay_other",
            "preheat_temperature",
            "interpass_temperature",
            "fillet_weld_position",
            "current",
            "voltage",
            "welding_speed",
            "heat_input",
            "filler_metal",
            "filler_diameter",
            "filler_standard",
            "filler_category",
            "filler_model",
            "filler_trade_name",
            "filler_class",
            "butt_weld_position",
            "vertical_direction",
            "pwht_temperature",
            "pwht_time",
            "polarity",
            "shielding_gas",
            "shielding_gas_mix",
            "gas_flow",
            "trailing_gas",
            "backing_gas",
            "preheat_time",
            "heating_method",
            "current_type",
            "tungsten_electrode",
            "nozzle_diameter",
            "arc_type",
            "wire_feed_speed",
            "weaving",
            "weaving_parameter",
            "cleaning",
            "back_gouging",
            "single_or_multi_pass",
            "single_or_multi_wire",
            "contact_tip_distance",
            "peening",
            "technical_other",
            "prepared_by",
            "prepared_date",
            "reviewed_by",
            "reviewed_date",
            "approved_by",
            "approved_date",
            "bead_1_process",
            "bead_1_filler_metal",
            "bead_1_diameter",
            "bead_1_polarity",
            "bead_1_current",
            "bead_1_voltage",
            "bead_1_speed",
            "bead_1_heat_input",
            "bead_2_process",
            "bead_2_filler_metal",
            "bead_2_diameter",
            "bead_2_polarity",
            "bead_2_current",
            "bead_2_voltage",
            "bead_2_speed",
            "bead_2_heat_input",
        )
        normalized = {
            key: WeldingDocumentAgent._clean_value(document_fields.get(key))
            for key in required_keys
        }
        for bead_no in (1, 2):
            normalized[f"bead_{bead_no}_process"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_process"),
                normalized.get("welding_process"),
            )
            normalized[f"bead_{bead_no}_filler_metal"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_filler_metal"),
                normalized.get("filler_trade_name"),
                normalized.get("filler_metal"),
            )
            normalized[f"bead_{bead_no}_diameter"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_diameter"),
                normalized.get("filler_diameter"),
            )
            normalized[f"bead_{bead_no}_polarity"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_polarity"),
                normalized.get("polarity"),
            )
            normalized[f"bead_{bead_no}_current"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_current"),
                normalized.get("current"),
            )
            normalized[f"bead_{bead_no}_voltage"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_voltage"),
                normalized.get("voltage"),
            )
            normalized[f"bead_{bead_no}_speed"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_speed"),
                normalized.get("welding_speed"),
            )
            normalized[f"bead_{bead_no}_heat_input"] = WeldingDocumentAgent._coalesce_value(
                normalized.get(f"bead_{bead_no}_heat_input"),
                normalized.get("heat_input"),
            )
        return normalized

    @staticmethod
    def _clean_value(value: Any) -> str:
        text = str(value).strip() if value is not None else ""
        lowered = text.lower()
        blocked = ("not found", "unknown tool", "unknown too", "error", "missing")
        if not text or any(item in lowered for item in blocked):
            return "/"
        return text

    @staticmethod
    def _coalesce_value(*values: Any) -> str:
        for value in values:
            clean_value = WeldingDocumentAgent._clean_value(value)
            if clean_value != "/":
                return clean_value
        return "/"

    @staticmethod
    def _value(document_fields: dict[str, str], key: str) -> str:
        return WeldingDocumentAgent._clean_value(document_fields.get(key))


def build_welding_document_agent(
    config: WeldingDocumentAgentConfig | None = None,
) -> WeldingDocumentAgent:
    return WeldingDocumentAgent(config=config)


def build_welding_document_agent_from_config(config: dict[str, Any]) -> WeldingDocumentAgent:
    template_config = config.get("template", {})
    output_config = config.get("output", {})
    return WeldingDocumentAgent(
        WeldingDocumentAgentConfig(
            template_docx_path=Path(template_config.get("docx_path", "data/MHPWPS-062.docx")),
            output_dir=Path(output_config.get("dir", "result")),
            output_prefix=output_config.get("prefix", "welding_standard_filled"),
        )
    )
